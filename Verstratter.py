from spacy.tokens import Doc
import spacy
import os
from docx import Document
import pandas as pd
import textract
from docxtpl import DocxTemplate
#pak de stratt template en vul hier alles op aan
docname = os.getcwd() + '\Stratt template.docx'
document = DocxTemplate(docname)
def print_doc_entities(_doc: Doc):
    if _doc.ents:
        for _ent in _doc.ents:
            print(f"     {_ent.text} {_ent.label_}")
    else:
        print("     NONE")
#Een cv van een kandidaat is gebruikt om de code te testen. Deze was in de "current working directory" opgeslagen als CV
txt = textract.process(os.getcwd() + "/CV.docx")
#decode de tekst in het cv
txt = txt.decode("utf-8")
#gebruik het model in de huidige schijf
nlp = spacy.load(os.getcwd())
#maak een doc container met entiteiten en label
doc = nlp(txt)
#print alle entiteiten en labels
print_doc_entities(doc)
counter = 0
periodeCounter = 0
#maak een dataframe van de entiteiten en labels. Hier kan logica op worden toegepast om de entiteiten die bij elkaar horen ook bij elkaar te zetten in het verstratte cv
df = pd.DataFrame(columns=["text","labels"])
#loop door alle entiteiten heen en voeg deze toe aan de dataframe
for _ent in doc.ents:
    print(_ent.text, _ent.label_)
    if _ent.label_ == "Periode":
        counter+=1
        df = df.append({'text':_ent.text, 'labels':_ent.label_}, ignore_index=True)
    if _ent.label_ == "Functie":
        counter+=1
        df = df.append({'text':_ent.text, 'labels':_ent.label_}, ignore_index=True)
    if _ent.label_ == "Werkgever":
        counter+=1
        df = df.append({'text':_ent.text, 'labels':_ent.label_}, ignore_index=True)
    if _ent.label_ == "Opleidingen":
        counter+=1
        df = df.append({'text':_ent.text, 'labels':_ent.label_}, ignore_index=True)
    if _ent.label_ == "Verantwoordelijkheden":
        counter+=1
        df = df.append({'text':_ent.text, 'labels':_ent.label_}, ignore_index=True)

print(df)
#maak een lijst met de posities in het dataframe van de periodes
periodeList = [df.loc[df['labels'] == 'Periode', 'text' ].index]
#variabele om de positie van de vorige index bij te houden
prevPeriod = 0
#maak kleinere dataframes voor alle bijeenhorende data.
for ExperienceCount in range(len(periodeList[0])):
    #temporary dataframe met bijbehorende data
    tempdf = df[prevPeriod:periodeList[0][ExperienceCount]]
    #sla een nieuwe previous periode op
    prevPeriod = periodeList[0][ExperienceCount]
    print(tempdf)
    tempLength = len(tempdf.index)
    print(tempLength)
#voegt de data toe aan het CV met behulp van python-docx. https://python-docx.readthedocs.io/en/latest/
#loop door alle doc entiteiten
for _ent in doc.ents:
    print(_ent.label_, _ent.text)
    addedOnce = False
    #loop door alle paragrafen heen van de template. Paragrafen moeten nog aangepast worden. tekst wordt soms toegevoegd aan het eind van opleidingen
    for p in document.paragraphs:
        if _ent.label_ == "Naam":
            if "Naam			:	" in p.text:
                #voeg de tekst toe aan het einde van de naam paragraaf
                p.add_run(_ent.text)
        if _ent.label_ == "Periode":
            #er staan 3 voorbeelden in het template van de werkervaring
            if periodeCounter < 3:
                # kijk door alle runs heen
                for run in p.runs:
                    #in het template staat x-x bij de periode
                    if run.text == "x-x":
                        #voeg dit maar 1 keer toe en zoek dan opnieuw naar x-x in het template
                        if addedOnce is False:
                            #voeg de entiteit tekst toe aan
                            run.text = _ent.text
                            addedOnce = True
                            periodeCounter += 1
            else:
                if addedOnce is False:
                    #voeg een nieuwe werkervaring aan het template toe
                    p.add_run("x-x	:	x bij x \n Verantwoordelijk voor: \n •	x \n •	x")
                    print("run toegevoegd")
                    periodeCounter += 1
                    for run in p.runs:
                        # in het template staat x-x bij de periode
                        if run.text == "x-x":
                            # voeg dit maar 1 keer toe en zoek dan opnieuw naar x-x in het template
                            if addedOnce is False:
                                # voeg de entiteit tekst toe aan
                                run.text = _ent.text
                                addedOnce = True
                                periodeCounter += 1